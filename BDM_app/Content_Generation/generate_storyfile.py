import os
import docx
import json

class Choice():
    def __init__(self, _id, description, res_node):
        self._id = _id
        self._description = description
        self._res_node = res_node

class Node():
    def __init__(self, _id, description, choices):
        self._id = _id
        self._description = description
        self._choices = choices

    def add_choice(self, choice):
        self._choices.append(choice)
    
    def extend_description(self, string):
        self._description += string

def find_node(nodes_list, title):
    '''
    Finds node with title in list
    '''
    for i, node in enumerate(nodes_list):
        if node._id == title:
            return i
    return None

def extract_nodes(input_path):
    '''
    from a txt/docx file, creates a list of nodes without adding the choices
    '''
    print(input_path)

    nodes = []
    processing_node = False

    '''doc = docx.Document(input_path)
    doc.save('test.docx')
    print(doc)
    _file = doc.paragraphs
    print(_file)
    '''
    with open(input_path, 'r', encoding='utf-8') as file:
    #with docx.Document(input_path).paragraphs as file:
        for line in file:
            # See if title
            if line.startswith('NAME = '):
                if processing_node:
                    nodes.append(current_node)
                title = line.split('NAME = ')[1]
                current_node = {'id': title, 'description': '', 'choices':[]}#Node(title, '', [])
                processing_node = True
                choice_count = 1

            if line.startswith('PAR = '):
                #current_node.extend_description(line.split('PAR = ')[1])
                current_node['description'] += line.split('PAR = ')[1]+'\n'

            if line.startswith('LINK = '):
                _id = choice_count
                description, res_node = line.split('LINK = ')[1].split('---')
                #current_node.add_choice(Choice(_id, description, res_node))
                current_node['choices'].append({'id': _id, 'description': description, 'res_node': res_node})
                choice_count += 1
            
    if processing_node:
        nodes.append(current_node)

    '''print(len(nodes))
    for node in nodes:
        print(len(node._choices))'''
    return nodes 


def parse_file(input_path, res_path):
    '''
    from a txt/docx file, creates a .js file with necessary data for the app (as a list of nodes)
    '''

    nodes = extract_nodes(input_path)

    with open(res_path, 'w') as outfile:
        json.dump(nodes, outfile)

if __name__=='__main__':
    cwd = os.getcwd()
    #extract_nodes(cwd+'/Prologue_Conclave_des_elfes_v1.docx')
    parse_file(cwd+'/test_text.txt', cwd+'/story2.json')